"""
Resume Parser
=============

Multi-format resume text extraction.
Supports PDF (via pdfplumber) and DOCX (via python-docx).
OCR fallback is deferred to Phase 7.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import BinaryIO, Union

import pdfplumber
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# Minimum characters to consider a parsed document "non-empty"
_MIN_TEXT_LENGTH = 50


# ── Data Classes ───────────────────────────────────────────────────────────

@dataclass
class ResumeDocument:
    """Container for parsed resume content."""

    raw_text: str
    file_type: str  # "pdf" | "docx" | "unknown"
    parsing_method: str  # "pdfplumber" | "python-docx" | "none"
    warnings: list[str] = field(default_factory=list)

    @property
    def is_empty(self) -> bool:
        return len(self.raw_text.strip()) < _MIN_TEXT_LENGTH

    @property
    def char_count(self) -> int:
        return len(self.raw_text)


# ── PDF Parsing ────────────────────────────────────────────────────────────

def parse_pdf(file: Union[str, Path, BinaryIO]) -> ResumeDocument:
    """
    Extract text from a PDF using pdfplumber.

    Parameters
    ----------
    file : str, Path, or file-like object
        Path to a PDF or an in-memory binary stream (e.g., from Streamlit upload).

    Returns
    -------
    ResumeDocument
    """
    warnings: list[str] = []
    text_parts: list[str] = []

    try:
        with pdfplumber.open(file) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                else:
                    warnings.append(f"Page {i + 1}: no text extracted (may be scanned/image-based)")
    except Exception as exc:
        logger.error("pdfplumber failed: %s", exc)
        return ResumeDocument(
            raw_text="",
            file_type="pdf",
            parsing_method="pdfplumber",
            warnings=[f"PDF parsing error: {exc}"],
        )

    raw_text = "\n\n".join(text_parts)
    doc = ResumeDocument(
        raw_text=raw_text,
        file_type="pdf",
        parsing_method="pdfplumber",
        warnings=warnings,
    )

    if doc.is_empty:
        doc.warnings.append(
            "Extracted text is very short or empty. "
            "This file may be a scanned document. OCR support is not yet enabled."
        )

    return doc


# ── DOCX Parsing ───────────────────────────────────────────────────────────

def parse_docx(file: Union[str, Path, BinaryIO]) -> ResumeDocument:
    """
    Extract text from a DOCX file using python-docx.

    Parameters
    ----------
    file : str, Path, or file-like object
        Path to a DOCX or an in-memory binary stream.

    Returns
    -------
    ResumeDocument
    """
    warnings: list[str] = []

    try:
        if isinstance(file, (str, Path)):
            doc = DocxDocument(str(file))
        else:
            doc = DocxDocument(file)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract text from tables (resumes often use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in paragraphs:
                        paragraphs.append(cell_text)

    except Exception as exc:
        logger.error("python-docx failed: %s", exc)
        return ResumeDocument(
            raw_text="",
            file_type="docx",
            parsing_method="python-docx",
            warnings=[f"DOCX parsing error: {exc}"],
        )

    raw_text = "\n".join(paragraphs)
    result = ResumeDocument(
        raw_text=raw_text,
        file_type="docx",
        parsing_method="python-docx",
        warnings=warnings,
    )

    if result.is_empty:
        result.warnings.append("Extracted text is very short or empty.")

    return result


# ── Auto-detect & Parse ───────────────────────────────────────────────────

def _detect_file_type(file: Union[str, Path, BinaryIO]) -> str:
    """Detect file type from extension or magic bytes."""
    if isinstance(file, (str, Path)):
        ext = Path(file).suffix.lower()
        if ext == ".pdf":
            return "pdf"
        elif ext in (".docx", ".doc"):
            return "docx"
        return "unknown"

    # For file-like objects, check the name attribute if available
    name = getattr(file, "name", "")
    if name:
        ext = Path(name).suffix.lower()
        if ext == ".pdf":
            return "pdf"
        elif ext in (".docx", ".doc"):
            return "docx"

    # Fallback: try to detect from content
    # Read first few bytes for magic number detection
    pos = file.tell() if hasattr(file, "tell") else 0
    header = file.read(8)
    file.seek(pos)  # Reset position

    if header.startswith(b"%PDF"):
        return "pdf"
    elif header.startswith(b"PK"):  # DOCX is a zip
        return "docx"

    return "unknown"


def parse_resume(file: Union[str, Path, BinaryIO]) -> ResumeDocument:
    """
    Parse a resume file, auto-detecting format.

    Supports PDF and DOCX. Falls back to an empty document with
    a warning for unsupported formats.

    Parameters
    ----------
    file : str, Path, or file-like object
        The resume file to parse.

    Returns
    -------
    ResumeDocument
    """
    file_type = _detect_file_type(file)

    if file_type == "pdf":
        return parse_pdf(file)
    elif file_type == "docx":
        return parse_docx(file)
    else:
        return ResumeDocument(
            raw_text="",
            file_type="unknown",
            parsing_method="none",
            warnings=["Unsupported file format. Please upload a PDF or DOCX file."],
        )
